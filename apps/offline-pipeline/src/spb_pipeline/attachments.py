from __future__ import annotations

import subprocess
from pathlib import Path
from urllib.parse import urlsplit

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from .io_utils import json_hash
from .models import Document
from .normalize import clean_text, validity_from_text


def _pdf_blocks(path: Path) -> tuple[list[dict], str]:
    reader = PdfReader(path)
    blocks: list[dict] = []
    for index, page in enumerate(reader.pages, 1):
        text = clean_text(page.extract_text() or "")
        if text:
            blocks.append({"type": "page", "page": index, "text": text})
    return blocks, "pdf_text" if blocks else "ocr_pending"


def _docx_blocks(path: Path) -> tuple[list[dict], str]:
    document = DocxDocument(path)
    blocks: list[dict] = []
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            blocks.append({"type": "paragraph", "text": text})
    for table in document.tables:
        rows = [
            [clean_text(cell.text) for cell in row.cells]
            for row in table.rows
        ]
        if rows:
            blocks.append({"type": "table", "headers": rows[0], "rows": rows[1:]})
    return blocks, "docx"


def _legacy_doc_blocks(path: Path) -> tuple[list[dict], str]:
    result = subprocess.run(
        ["textutil", "-convert", "txt", "-stdout", str(path)],
        check=True,
        capture_output=True,
    )
    text = clean_text(result.stdout.decode("utf-8", errors="replace"))
    blocks = [
        {"type": "paragraph", "text": line}
        for line in text.splitlines()
        if clean_text(line)
    ]
    return blocks, "legacy_doc"


def _xlsx_blocks(path: Path) -> tuple[list[dict], str]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[dict] = []
    for sheet in workbook.worksheets:
        rows = [
            [clean_text("" if value is None else str(value)) for value in row]
            for row in sheet.iter_rows(values_only=True)
        ]
        rows = [row for row in rows if any(row)]
        if rows:
            blocks.append(
                {
                    "type": "table",
                    "sheet": sheet.title,
                    "headers": rows[0],
                    "rows": rows[1:],
                }
            )
    return blocks, "xlsx"


def _blocks_to_text(blocks: list[dict]) -> str:
    values: list[str] = []
    for block in blocks:
        if block["type"] == "table":
            if block.get("headers"):
                values.append(" | ".join(block["headers"]))
            values.extend(" | ".join(row) for row in block.get("rows") or [])
        else:
            values.append(block.get("text") or "")
    return clean_text("\n\n".join(values))


def parse_attachment(
    attachment: dict,
    *,
    local_path: Path | None,
    parent_published_at: str,
    parent_document_no: str,
    parent_source_org: str,
    ocr_path: Path | None = None,
) -> Document:
    suffix = (
        local_path.suffix.lower()
        if local_path
        else Path(urlsplit(attachment["source_url"]).path).suffix.lower()
    )
    blocks: list[dict] = []
    parse_method = "not_downloaded"
    fetch_status = attachment.get("fetch_status") or "pending"
    if (
        local_path
        and local_path.exists()
        and local_path.stat().st_size > 0
        and fetch_status == "success"
    ):
        try:
            if suffix == ".pdf":
                blocks, parse_method = _pdf_blocks(local_path)
                if (
                    not blocks
                    and ocr_path
                    and ocr_path.exists()
                ):
                    from .io_utils import read_jsonl

                    blocks = list(read_jsonl(ocr_path))
                    parse_method = "vision_ocr"
            elif suffix == ".docx":
                if local_path.read_bytes()[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
                    blocks, parse_method = _legacy_doc_blocks(local_path)
                else:
                    blocks, parse_method = _docx_blocks(local_path)
            elif suffix in {".xlsx", ".xlsm"}:
                blocks, parse_method = _xlsx_blocks(local_path)
            elif suffix == ".txt":
                blocks = [
                    {
                        "type": "paragraph",
                        "text": clean_text(
                            local_path.read_text(
                                encoding="utf-8", errors="replace"
                            )
                        ),
                    }
                ]
                parse_method = "text"
            elif suffix in {".jpg", ".jpeg", ".png", ".tif", ".tiff"}:
                if ocr_path and ocr_path.exists():
                    from .io_utils import read_jsonl

                    blocks = list(read_jsonl(ocr_path))
                    parse_method = "vision_ocr"
                else:
                    parse_method = "ocr_pending"
            elif suffix == ".doc":
                blocks, parse_method = _legacy_doc_blocks(local_path)
            else:
                parse_method = "unsupported"
            fetch_status = "success"
        except Exception as exc:  # 保留单个坏附件，不中断全量解析。
            parse_method = "parse_failed"
            fetch_status = "failed"
            blocks = [{"type": "error", "text": str(exc)}]
    text = _blocks_to_text([block for block in blocks if block["type"] != "error"])
    if parse_method in {"ocr_pending", "conversion_pending"}:
        fetch_status = parse_method
    return Document(
        document_id=attachment["attachment_id"],
        parent_document_id=attachment["parent_document_id"],
        title=attachment["title"],
        source_url=attachment["source_url"],
        source_host=(urlsplit(attachment["source_url"]).hostname or "").lower(),
        document_type=suffix.lstrip(".") or "attachment",
        published_at=parent_published_at,
        document_no=parent_document_no,
        source_org=parent_source_org,
        validity_status=validity_from_text(attachment["title"], text),
        template_type="attachment",
        blocks=blocks,
        text=text,
        attachments=[],
        related_links=[],
        content_hash=json_hash({"blocks": blocks, "source_url": attachment["source_url"]}),
        fetch_status=fetch_status if text or fetch_status != "success" else "empty",
        parse_method=parse_method,
        metadata={
            "relation": attachment.get("relation", "attachment"),
            "discovered_from": attachment.get("discovered_from", ""),
            "local_path": str(local_path) if local_path else "",
        },
    )
