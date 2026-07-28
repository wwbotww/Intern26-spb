from __future__ import annotations

from docx import Document as DocxDocument
from openpyxl import Workbook

from spb_pipeline.attachments import parse_attachment


def _attachment(source_url: str) -> dict:
    return {
        "attachment_id": "attachment-1",
        "parent_document_id": "parent-1",
        "title": "政策附件",
        "source_url": source_url,
        "relation": "attachment",
        "discovered_from": "html",
        "fetch_status": "success",
    }


def test_parse_docx_attachment(tmp_path):
    path = tmp_path / "policy.docx"
    source = DocxDocument()
    source.add_heading("附件标题", level=1)
    source.add_paragraph("第一条 附件正文。")
    source.save(path)

    document = parse_attachment(
        _attachment("https://example.test/policy.docx"),
        local_path=path,
        parent_published_at="2026-01-01",
        parent_document_no="国邮发〔2026〕1号",
        parent_source_org="国家邮政局",
    )

    assert document.parse_method == "docx"
    assert document.fetch_status == "success"
    assert "第一条 附件正文" in document.text


def test_parse_xlsx_attachment(tmp_path):
    path = tmp_path / "standards.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "标准目录"
    sheet.append(["标准号", "标准名称"])
    sheet.append(["YZ/T 0001", "测试标准"])
    workbook.save(path)

    document = parse_attachment(
        _attachment("https://example.test/standards.xlsx"),
        local_path=path,
        parent_published_at="2026-01-01",
        parent_document_no="",
        parent_source_org="国家邮政局",
    )

    assert document.parse_method == "xlsx"
    assert document.fetch_status == "success"
    assert "YZ/T 0001 | 测试标准" in document.text
